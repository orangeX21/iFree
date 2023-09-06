import docx

doc = docx.Document('name.docx')

kept_paras = []
#
# import docx
#
# doc = docx.Document('file.docx')
# kept_paras = []

for para in doc.paragraphs:
    if '-' in para.text:
        kept_paras.append(para)

# 后续保存到新文件......

new_doc = docx.Document()
for para in kept_paras:
    new_doc.add_paragraph(para.text)

# new_doc.save('new_file.docx')