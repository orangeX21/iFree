import csv

with open('date.csv', 'r') as f:
    reader = csv.reader(f)
    header = next(reader)
    first_row = next(reader)

import docx

doc = docx.Document('new_file.docx')
# for para in doc.paragraphs:
#     print(para.text)

with open('output.txt', 'w') as f:
    for i in range(len(header)):
        # for para in doc.paragraphs:
        #     print(para.text)
        # print(header[i])
        f.write(header[i] + ':' + first_row[i] + "\n")
txt_dict = {}

with open('output.txt') as f:
    for line in f:
        field, name = line.split(':')
        txt_dict[field] = name
# print(txt_dict)
# output.txt解析到字典
# out_dict = {}

# docx解析到字典
doc_dict = {}
import docx

doc = docx.Document('new_file.docx')

# doc_dict = {}

import re

# 编译正则模式
pattern = r'(\w+) - (.*)'

# 遍历段落
for para in doc.paragraphs:

    # 使用正则匹配
    m = re.match(pattern, para.text)

    if m:
        # 从匹配结果提取字段和值
        field = m.group(1)
        value = m.group(2)

        # 保存到字典
        doc_dict[field] = value
# print(doc_dict)
# print(doc_dict)
# 逐字段匹配和合并
# 匹配和合并
test = []
with open('put.txt', 'w') as f:
    for k in txt_dict:

        # if k in doc_dict:
        #     v = doc_dict[k] + '(' + txt_dict[k] + ')'
        #
        # else:
        #     v = txt_dict[k]
        #
        # f.write(k + ' - ' + "".join(v).replace('\n','')+':' +"\n"+"\n")
        # print(k)

        # ***************
        name = k
        print(name)
        # num = input(k+":")
        # if num != "0":
        #     print("""dataset <- dataset %>%
        #           mutate_at(vars({}),
        #                     list(~ ifelse(. %in% c({}), NA, .)))""".format(name, num),end="\n")
        # else:
        #     pass
        # print("""filter(!is.na({})) %>%""".format(name))
# while 1:
#     name = input("name:")
#     print("""{} = as.factor({})""".format(name,name))